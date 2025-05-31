import requests
from PIL import Image
from io import BytesIO
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

def main():
    # Demander URL de la page HTML
    url_page = input("Entrez l'URL de la page HTML (ex: https://www.gutenberg.org/cache/epub/1342/pg1342-images.html) : ").strip()

    # Télécharger la page HTML
    response = requests.get(url_page)
    response.raise_for_status()
    html = response.text

    # Parser la page HTML pour extraire le titre et l'auteur
    soup = BeautifulSoup(html, 'html.parser')
    title_tag = soup.find('title')
    if title_tag:
        title_text = title_tag.text
        # Gutenberg: titre et auteur sont souvent dans le <title> comme "Pride and Prejudice by Jane Austen"
        if " by " in title_text:
            title, author = map(str.strip, title_text.split(" by ", 1))
        else:
            title = title_text.strip()
            author = "Auteur inconnu"
    else:
        title = "Titre inconnu"
        author = "Auteur inconnu"

    # Chemin relatif de l'image (à adapter si nécessaire)
    img_relative_url = "images/cover.jpg"
    img_url = urljoin(url_page, img_relative_url)
    print("URL complète de l'image :", img_url)

    # Télécharger l'image
    response_img = requests.get(img_url)
    response_img.raise_for_status()

    # Ouvrir et sauvegarder l'image localement (temporaire)
    img = Image.open(BytesIO(response_img.content))
    img_path = "cover_image.jpg"
    img.save(img_path)
    print(f"Image sauvegardée dans '{img_path}'")

    # Demander l'auteur du rapport
    report_author = input("Entrez le nom de l'auteur du rapport : ").strip()

    # Création du document Word
    doc = Document()

    # Page de titre
    doc.add_heading("Rapport de Lecture", level=0)
    doc.add_paragraph(f"Titre du livre : {title}")
    doc.add_paragraph(f"Auteur du livre : {author}")
    doc.add_paragraph(f"Auteur du rapport : {report_author}")
    doc.add_picture(img_path, width=Inches(4))

    # Sauvegarder le document
    word_path = "rapport.docx"
    doc.save(word_path)
    print(f"Document Word sauvegardé sous '{word_path}'")

if __name__ == "__main__":
    main()

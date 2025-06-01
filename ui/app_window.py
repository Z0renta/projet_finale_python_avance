import tkinter as tk
from tkinter import messagebox, simpledialog
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from database import create_table, clear_table, insert_data, fetch_all
from api import fetch_posts
import os
import requests
import re
from collections import Counter
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
from collections import defaultdict
from bs4 import BeautifulSoup
from urllib.parse import urljoin


class Application(tk.Tk):
    def __init__(self):
        self.canvas = None

        super().__init__()
        self.title("Application de Données JSON")
        self.geometry("800x600")
        create_table()
        self.create_menu()
        self.create_widgets()

    def create_menu(self):
        menu_bar = tk.Menu(self)

        data_menu = tk.Menu(menu_bar, tearoff=0)
        data_menu.add_command(label="Effacer la base de données", command=self.clear_database)
        data_menu.add_command(label="Télécharger les données", command=self.download_data)
        data_menu.add_command(label="Afficher le graphique", command=self.show_graph)
        data_menu.add_command(label="Afficher l'agrégation", command=self.show_aggregation)
        menu_bar.add_cascade(label="Données", menu=data_menu)

        book_menu = tk.Menu(menu_bar, tearoff=0)
        book_menu.add_command(label="Traiter un livre", command=self.process_book)
        menu_bar.add_cascade(label="Livre", menu=book_menu)

        options_menu = tk.Menu(menu_bar, tearoff=0)
        options_menu.add_command(label="Changer la couleur", command=self.set_bg_color)
        menu_bar.add_cascade(label="Options", menu=options_menu)

        self.config(menu=menu_bar)

    def create_widgets(self):
        self.text_area = tk.Text(self)
        self.text_area.pack(fill=tk.BOTH, expand=True)

    def clear_database(self):
        confirm = messagebox.askyesno("Confirmation", "Supprimer complètement le fichier de base de données ?")
        if confirm:
            db_path = "data/donnees.db"
            if os.path.exists(db_path):
                # os.remove(db_path)
                clear_table()
                self.destroy_graph()
                self.canvas = None
                self.text_area.delete(1.0, tk.END)
                messagebox.showinfo("Succès", "Le fichier de base de données a été supprimé.")
            else:
                messagebox.showwarning("Erreur", "Le fichier n'existe pas.")

    def download_data(self):
        try:
            data = fetch_posts()
            insert_data(data)
            self.display_data()
        except Exception as e:
            messagebox.showerror("Erreur", str(e))

    def display_data(self):
        rows = fetch_all()
        self.text_area.delete(1.0, tk.END)
        for row in rows:
            self.text_area.insert(tk.END, f"ID: {row[0]}, UserID: {row[1]}\nTitre: {row[2]}\nCorps: {row[3]}\n\n")

    def show_graph(self):
        rows = fetch_all()
        if rows != []:
            user_counts = {}
            for row in rows:
                user_counts[row[1]] = user_counts.get(row[1], 0) + 1

            fig, ax = plt.subplots()
            ax.bar(user_counts.keys(), user_counts.values())
            ax.set_xlabel("UserID")
            ax.set_ylabel("Nombre de posts")
            ax.set_title("Posts par utilisateur")

            self.canvas = FigureCanvasTkAgg(fig, master=self)
            self.canvas.draw()
            self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
            self.canvas_widget = self.canvas.get_tk_widget()
            self.canvas_widget.pack(fill=tk.BOTH, expand=True)
        else:
            messagebox.showerror("Attention", "Vous devez avant tout télécharger les données !")
    
    def destroy_graph(self):
        if hasattr(self, 'canvas_widget') and self.canvas_widget:
            self.canvas_widget.destroy()
            self.canvas = None
            self.canvas_widget = None


    def show_aggregation(self):
        rows = fetch_all()
        if not rows:
            messagebox.showinfo("information", "Aucune donnée à afficher.")
            return

        user_length = defaultdict(list)

        for row in rows:
            user_id = row[1]
            body_length = len(row[3])
            user_length[user_id].append(body_length)

        user_avg = {user: sum(lengths) / len(lengths) for user, lengths in user_length.items()}

        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, "Moyenne des longueurs de texte par utilisateur :\n\n")
        for user, avg in sorted(user_avg.items()):
            self.text_area.insert(tk.END, f"User {user} : {avg:.2f} caractères en moyenne\n")


        fig, ax = plt.subplots()
        ax.bar(user_avg.keys(), user_avg.values())
        ax.set_title("Longueur moyenne des posts par utilisateur")
        ax.set_xlabel("User ID")
        ax.set_ylabel("Longueur moyenne (caractères)")

        self.canvas = FigureCanvasTkAgg(fig, master=self)
        self.canvas.draw()
        self.canvas_widget = self.canvas.get_tk_widget()
        self.canvas_widget.pack(fill=tk.BOTH, expand=True)


    def set_bg_color(self):
        list_color = ['red', 'blue', 'green']
        boucle = True
        while boucle:
            color = simpledialog.askstring("Options", "Entrez une couleur de fond (red, green, blue):")
            if color in list_color:
                boucle = False
                self.text_area.config(bg=color)
            else :
                messagebox.showinfo("Erreur", "Vous devez entrer une couleur de la liste !")


 # ======= MÉTHODES LIVRE =======

    def download_book_text(self, url):
        response = requests.get(url)
        response.encoding = 'utf-8'
        return response.text

    def extract_metadata_and_first_chapter(self, text):
        title_match = re.search(r'Title:\s*(.*)', text)
        title = title_match.group(1).strip() if title_match else "Inconnu"

        author_match = re.search(r'Author:\s*(.*)', text)
        author = author_match.group(1).strip() if author_match else "Inconnu"

        chapters = re.split(r'\nChapter [0-9IVXLC]+\.*\s*\n', text, flags=re.IGNORECASE)
        if len(chapters) > 1:
            first_chapter = chapters[1].strip()
        else:
            first_chapter = text

        return title, author, first_chapter

    def count_words_per_paragraph(self, chapter_text):
        paragraphs = [p.strip() for p in chapter_text.split('\n\n') if p.strip()]
        counts = []
        for p in paragraphs:
            nb_mots = len(p.split())
            nb_mots_rounded = round(nb_mots / 10) * 10
            counts.append(nb_mots_rounded)
        return counts

    def paragraph_length_distribution(self, counts):
        counter = Counter(counts)
        sorted_counts = sorted(counter.items())
        return sorted_counts


    def crop_and_resize_image(self, path, crop_box=(100, 100, 400, 400), size=(300, 300)):
        img = Image.open(path)
        cropped = img.crop(crop_box)
        resized = cropped.resize(size)
        resized.save(path)


    def process_book(self):
        url_page = simpledialog.askstring("Livre", "Entrez l'URL de la page HTML (ex: https://www.gutenberg.org/cache/epub/1342/pg1342-images.html) :")
        if not url_page:
            return

        try:
            # Télécharger la page HTML
            headers = {
                "User-Agent": "Mozilla/5.0"
            }
            response = requests.get(url_page, headers=headers)
            response.raise_for_status()
            html = response.text

            # Parser avec BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            title_tag = soup.find('title')
            if title_tag and " by " in title_tag.text:
                title, author = map(str.strip, title_tag.text.split(" by ", 1))
            else:
                title = title_tag.text.strip() if title_tag else "Titre inconnu"
                author = "Auteur inconnu"

            # Construire URL de l'image (image par défaut : cover.jpg dans le dossier "images/")
            img_relative_url = "images/cover.jpg"
            img_url = urljoin(url_page, img_relative_url)

            # Télécharger l'image avec header User-Agent
            img_response = requests.get(img_url, headers=headers)
            if img_response.status_code != 200:
                raise Exception(f"Erreur lors du téléchargement de l'image : {img_response.status_code}")

            img = Image.open(BytesIO(img_response.content))
            img_path = "cover_image.jpg"
            

            img.save(img_path)

            # Demander auteur du rapport
            report_author = simpledialog.askstring("Auteur du rapport", "Entrez le nom de l'auteur du rapport :")
            if not report_author:
                report_author = "Inconnu"

            # Créer le document Word
            doc = Document()
            doc.add_heading("Rapport de Lecture", level=0)
            doc.add_paragraph(f"Titre du livre : {title}")
            doc.add_paragraph(f"Auteur du livre : {author}")
            doc.add_paragraph(f"Auteur du rapport : {report_author}")
            doc.add_picture(img_path, width=Inches(4))

            doc.save("rapport.docx")
            messagebox.showinfo("Succès", "Document Word généré : rapport.docx")

            # Afficher infos dans la zone de texte
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(tk.END, f"📘 Livre traité : {title} par {author}\nDocument Word sauvegardé.")
            # Télécharger le texte brut du livre
            txt_url = url_page.replace("-images.html", ".txt")
            text = self.download_book_text(txt_url)
            _, _, first_chapter = self.extract_metadata_and_first_chapter(text)

            # Compter les mots par paragraphe
            counts = self.count_words_per_paragraph(first_chapter)
            distribution = self.paragraph_length_distribution(counts)

            # Afficher la distribution dans un graphique
            x_vals, y_vals = zip(*distribution)
            fig, ax = plt.subplots()
            ax.bar(x_vals, y_vals, color='skyblue')
            ax.set_xlabel("Longueur du paragraphe (arrondi à la dizaine)")
            ax.set_ylabel("Nombre de paragraphes")
            ax.set_title("Distribution des longueurs de paragraphes (Chapitre 1)")
            self.destroy_graph()  # Nettoyer les anciens graphes
            self.canvas = FigureCanvasTkAgg(fig, master=self)
            self.canvas.draw()
            self.canvas_widget = self.canvas.get_tk_widget()
            self.canvas_widget.pack(fill=tk.BOTH, expand=True)


        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur est survenue : {str(e)}")


    def on_close(self):
        print("Fermeture interceptée.")
        # Tu peux afficher un message, enregistrer des données, etc.
        if tk.messagebox.askokcancel("Quitter", "Veux-tu vraiment quitter ?"):
            clear_table()
            self.destroy()